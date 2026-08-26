"""Document parsers and chunkers for PDF, DOCX, and CSV HR documents."""

from __future__ import annotations

import csv
import hashlib
import io
import re
from pathlib import Path
from typing import Sequence

import docx
import pdfplumber

from hr_chatbot.domain import DocumentVersion, KnowledgeChunk


def compute_file_hash(path: Path) -> str:
    """Compute SHA-256 hash of a file."""
    hasher = hashlib.sha256()
    with open(path, "rb") as f:
        while chunk := f.read(65536):
            hasher.update(chunk)
    return hasher.hexdigest()


def make_chunk_id(doc_id: str, section: str, text: str) -> str:
    """Generate a stable, deterministic chunk ID."""
    raw = f"{doc_id}:{section}:{text.strip()}"
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()[:16]


class DocumentParser:
    """Multi-format parser for HR policy documents."""

    def parse_file(self, file_path: Path) -> tuple[DocumentVersion, list[KnowledgeChunk]]:
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            return self.parse_pdf(file_path)
        elif suffix == ".docx":
            return self.parse_docx(file_path)
        elif suffix in (".csv", ".tsv"):
            return self.parse_csv(file_path)
        else:
            return self.parse_text(file_path)

    def parse_pdf(self, path: Path) -> tuple[DocumentVersion, list[KnowledgeChunk]]:
        doc_id = path.stem.replace(" ", "_").lower()
        title = path.stem
        file_hash = compute_file_hash(path)
        doc_ver = DocumentVersion(
            document_id=doc_id,
            version_id="v1",
            title=title,
            document_kind="rule",
            source_uri=str(path),
            priority=300,
            content_hash=file_hash,
        )

        chunks: list[KnowledgeChunk] = []
        with pdfplumber.open(path) as pdf:
            current_section = "일반"
            buffer: list[str] = []
            page_num = 1

            for p_idx, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text() or ""
                lines = page_text.split("\n")

                for line in lines:
                    line_str = line.strip()
                    if not line_str:
                        continue

                    # Detect section / article patterns like "제 10 조(근무시간)" or "제1장 총칙"
                    match = re.match(r"^(제\s*\d+\s*[조장절관]\s*(\([^\)]+\))?)", line_str)
                    if match:
                        if buffer:
                            chunk_text = "\n".join(buffer).strip()
                            if len(chunk_text) > 20:
                                search_text = f"[{title}] {current_section}\n{chunk_text}"
                                chunks.append(
                                    KnowledgeChunk(
                                        chunk_id=make_chunk_id(doc_id, current_section, chunk_text),
                                        document_id=doc_id,
                                        version_id="v1",
                                        title=title,
                                        page_or_section=f"p.{page_num} {current_section}",
                                        text=chunk_text,
                                        search_text=search_text,
                                        document_kind="rule",
                                        priority=300,
                                        source_uri=str(path),
                                    )
                                )
                            buffer = []
                        current_section = line_str[:40]
                        page_num = p_idx

                    buffer.append(line_str)

                # Extract tables if any
                tables = page.extract_tables() or []
                for t_idx, table in enumerate(tables):
                    if not table:
                        continue
                    headers = [str(c or "").strip() for c in table[0] if c]
                    header_str = " | ".join(headers)
                    table_rows: list[str] = []
                    for row in table[1:]:
                        clean_cells = [str(c or "").strip() for c in row if c is not None]
                        if any(clean_cells):
                            table_rows.append(" | ".join(clean_cells))
                    if table_rows:
                        table_text = f"표 (헤더: {header_str})\n" + "\n".join(table_rows)
                        sec = f"p.{p_idx} 표 {t_idx+1} ({current_section})"
                        chunks.append(
                            KnowledgeChunk(
                                chunk_id=make_chunk_id(doc_id, sec, table_text),
                                document_id=doc_id,
                                version_id="v1",
                                title=title,
                                page_or_section=sec,
                                text=table_text,
                                search_text=f"[{title}] {sec}\n{table_text}",
                                document_kind="rule",
                                priority=300,
                                source_uri=str(path),
                                table_id=f"tab_{p_idx}_{t_idx}",
                            )
                        )

            if buffer:
                chunk_text = "\n".join(buffer).strip()
                if len(chunk_text) > 20:
                    search_text = f"[{title}] {current_section}\n{chunk_text}"
                    chunks.append(
                        KnowledgeChunk(
                            chunk_id=make_chunk_id(doc_id, current_section, chunk_text),
                            document_id=doc_id,
                            version_id="v1",
                            title=title,
                            page_or_section=f"p.{page_num} {current_section}",
                            text=chunk_text,
                            search_text=search_text,
                            document_kind="rule",
                            priority=300,
                            source_uri=str(path),
                        )
                    )

        return doc_ver, chunks

    def parse_docx(self, path: Path) -> tuple[DocumentVersion, list[KnowledgeChunk]]:
        doc_id = path.stem.replace(" ", "_").lower()
        title = path.stem
        file_hash = compute_file_hash(path)
        doc_ver = DocumentVersion(
            document_id=doc_id,
            version_id="v1",
            title=title,
            document_kind="rule",
            source_uri=str(path),
            priority=300,
            content_hash=file_hash,
        )

        doc = docx.Document(path)
        chunks: list[KnowledgeChunk] = []
        current_section = "일반"
        buffer: list[str] = []

        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue

            match = re.match(r"^(제\s*\d+\s*[조장절관]\s*(\([^\)]+\))?)", text)
            if match:
                if buffer:
                    chunk_text = "\n".join(buffer).strip()
                    if len(chunk_text) > 20:
                        search_text = f"[{title}] {current_section}\n{chunk_text}"
                        chunks.append(
                            KnowledgeChunk(
                                chunk_id=make_chunk_id(doc_id, current_section, chunk_text),
                                document_id=doc_id,
                                version_id="v1",
                                title=title,
                                page_or_section=current_section,
                                text=chunk_text,
                                search_text=search_text,
                                document_kind="rule",
                                priority=300,
                                source_uri=str(path),
                            )
                        )
                    buffer = []
                current_section = text[:40]

            buffer.append(text)

        for t_idx, table in enumerate(doc.tables):
            table_rows: list[str] = []
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    table_rows.append(" | ".join(cells))
            if table_rows:
                table_text = "\n".join(table_rows)
                sec = f"표 {t_idx+1}"
                chunks.append(
                    KnowledgeChunk(
                        chunk_id=make_chunk_id(doc_id, sec, table_text),
                        document_id=doc_id,
                        version_id="v1",
                        title=title,
                        page_or_section=sec,
                        text=table_text,
                        search_text=f"[{title}] {sec}\n{table_text}",
                        document_kind="rule",
                        priority=300,
                        source_uri=str(path),
                        table_id=f"tab_{t_idx}",
                    )
                )

        if buffer:
            chunk_text = "\n".join(buffer).strip()
            if len(chunk_text) > 20:
                search_text = f"[{title}] {current_section}\n{chunk_text}"
                chunks.append(
                    KnowledgeChunk(
                        chunk_id=make_chunk_id(doc_id, current_section, chunk_text),
                        document_id=doc_id,
                        version_id="v1",
                        title=title,
                        page_or_section=current_section,
                        text=chunk_text,
                        search_text=search_text,
                        document_kind="rule",
                        priority=300,
                        source_uri=str(path),
                    )
                )

        return doc_ver, chunks

    def parse_csv(self, path: Path) -> tuple[DocumentVersion, list[KnowledgeChunk]]:
        doc_id = path.stem.replace(" ", "_").lower()
        title = path.stem
        file_hash = compute_file_hash(path)
        doc_ver = DocumentVersion(
            document_id=doc_id,
            version_id="v1",
            title=title,
            document_kind="faq",
            source_uri=str(path),
            priority=200,
            content_hash=file_hash,
        )

        chunks: list[KnowledgeChunk] = []
        with open(path, "r", encoding="utf-8-sig") as f:
            reader = csv.DictReader(f)
            for idx, row in enumerate(reader, start=1):
                category = row.get("대분류", "")
                sub_cat = row.get("소분류", "")
                q_or_sec = row.get("질문/조항", row.get("질문", row.get("항목", f"항목 {idx}")))
                content = row.get("내용/답변", row.get("답변", row.get("내용", "")))
                keywords = row.get("활용_키워드", "")
                gubun = row.get("구분", "FAQ")

                section_label = f"[{category} > {sub_cat}] {q_or_sec}" if category else q_or_sec
                search_text = f"[{title}] {section_label}\n{content}\n키워드: {keywords}"

                chunks.append(
                    KnowledgeChunk(
                        chunk_id=make_chunk_id(doc_id, section_label, content),
                        document_id=doc_id,
                        version_id="v1",
                        title=f"{title} ({gubun})",
                        page_or_section=section_label,
                        text=content,
                        search_text=search_text,
                        document_kind="faq" if "FAQ" in gubun else "rule",
                        priority=150 if "FAQ" in gubun else 250,
                        source_uri=str(path),
                    )
                )

        return doc_ver, chunks

    def parse_text(self, path: Path) -> tuple[DocumentVersion, list[KnowledgeChunk]]:
        doc_id = path.stem.replace(" ", "_").lower()
        title = path.stem
        file_hash = compute_file_hash(path)
        doc_ver = DocumentVersion(
            document_id=doc_id,
            version_id="v1",
            title=title,
            document_kind="rule",
            source_uri=str(path),
            priority=100,
            content_hash=file_hash,
        )

        content = path.read_text(encoding="utf-8")
        paragraphs = [p.strip() for p in content.split("\n\n") if len(p.strip()) > 20]
        chunks: list[KnowledgeChunk] = []

        for idx, para in enumerate(paragraphs, start=1):
            sec = f"단락 {idx}"
            chunks.append(
                KnowledgeChunk(
                    chunk_id=make_chunk_id(doc_id, sec, para),
                    document_id=doc_id,
                    version_id="v1",
                    title=title,
                    page_or_section=sec,
                    text=para,
                    search_text=f"[{title}] {sec}\n{para}",
                    document_kind="rule",
                    priority=100,
                    source_uri=str(path),
                )
            )

        return doc_ver, chunks
